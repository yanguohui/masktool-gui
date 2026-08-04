"""
训练「通信工程领域」spaCy NER 模型
==================================

**为什么需要自己训练？**

公开渠道并不存在针对通信工程领域微调的中文 spaCy NER 模型。
spaCy 官方的 ``zh_core_web_*`` 基于 OntoNotes 5 通用语料训练，在技术
规范书 / 招标文件上有两个硬伤：实体常被截断（「昆山市轨道交通集团
有限公司」只识别出「交通集团有限公司」），且完全不认识「标段」「网元」
「承载网」这类领域专名。因此**必须用自己的文档训练**。

**怎么解决「没有标注数据」的问题？——弱监督自动标注**

本脚本不要求人工标注。它复用本项目已有的**构词法识别层**
（``app.engine`` 里那套组织形式后缀 / 行政区划 / 项目中心词 / 词性定界
规则）来自动给语料打标签，再用这批标签去训练 spaCy。

    真实文档 ──构词法规则──▶ 自动标注 ──▶ 训练 spaCy ──▶ 领域模型

这样得到的模型能把规则「蒸馏」进统计模型，从而具备规则写不出来的
泛化能力（比如认出没有标准后缀的简称）。规则层依旧保留，两者互为补充。

用法
----
1) 准备语料：把若干真实文档（.docx/.txt）放进一个目录，例如 ``corpus/``
   注意：训练只在本机进行，语料不会外传。

2) 生成训练集并训练::

       python tools/train_spacy_ner.py --corpus corpus/ --out models/telecom_ner

3) 训练完成后程序会自动发现 ``models/telecom_ner``
   （目录名含 ``telecom`` 会被优先加载），在界面里把 NER 后端设为
   ``auto`` 或 ``spacy`` 即可生效。

依赖::

    pip install spacy
    python -m spacy download zh_core_web_sm    # 作为训练起点（可选但推荐）
"""

from __future__ import annotations

import argparse
import random
import sys
from pathlib import Path

# 让脚本能直接从项目根目录运行
_ROOT = Path(__file__).resolve().parent.parent
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

#: mask-tool 实体类别 → spaCy 标签
TYPE_TO_LABEL = {
    "company": "ORG",
    "government": "ORG",
    "person": "PERSON",
    "location": "GPE",
    "project": "PROJECT",
}


# --------------------------------------------------------------------------
# 语料读取
# --------------------------------------------------------------------------


def read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        return ""
    try:
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(c.text for c in row.cells)
        return "\n".join(parts)
    except Exception:
        return ""


def read_text(path: Path) -> str:
    for enc in ("utf-8", "gbk"):
        try:
            return path.read_text(encoding=enc)
        except (OSError, UnicodeDecodeError, LookupError):
            continue
    return ""


def load_corpus(corpus_dir: Path) -> list[str]:
    """读取语料目录，返回段落列表（按空行/换行切段，过滤过短的行）。"""
    texts: list[str] = []
    for p in sorted(corpus_dir.rglob("*")):
        if not p.is_file():
            continue
        ext = p.suffix.lower()
        if ext == ".docx":
            raw = read_docx(p)
        elif ext in (".txt", ".md"):
            raw = read_text(p)
        else:
            continue
        if raw:
            texts.append(raw)

    chunks: list[str] = []
    for raw in texts:
        for line in raw.split("\n"):
            s = line.strip()
            if len(s) >= 8:
                chunks.append(s)
    return chunks


# --------------------------------------------------------------------------
# 弱监督自动标注
# --------------------------------------------------------------------------


def auto_annotate(chunks: list[str]) -> list[tuple[str, dict]]:
    """用本项目的构词法规则给语料自动打标签。

    Returns:
        [(文本, {"entities": [(start, end, label), ...]}), ...]
    """
    sys.path.insert(0, str(_ROOT))
    import app.engine as E

    # 触发规则补丁（把自主识别正则注入 mask-tool 的 Detector）
    E.MaskEngine(E.locate_mask_tool())

    from mask_tool.core.detector import Detector
    from mask_tool.models.config import MaskConfig

    cfg = MaskConfig(mode="smart")
    wl = set(E.load_whitelist(force=True))
    det = Detector({}, wl, ner_engine=None)   # 只用正则层，不用 NER

    data: list[tuple[str, dict]] = []
    for text in chunks:
        try:
            results = det.detect(text, file_path="corpus")
        except Exception:
            continue

        spans: list[tuple[int, int, str]] = []
        for r in results:
            if getattr(r, "source", "") != "regex":
                continue
            if float(getattr(r, "confidence", 0)) < 0.85:
                continue
            name = getattr(r.text_type, "value", str(r.text_type)).lower()
            label = TYPE_TO_LABEL.get(name)
            if label is None:
                continue
            if E.is_whitelisted(r.text):
                continue
            idx = text.find(r.text)
            if idx < 0:
                continue
            spans.append((idx, idx + len(r.text), label))

        spans = _drop_overlaps(spans)
        if spans:
            data.append((text, {"entities": spans}))
    return data


def _drop_overlaps(spans: list[tuple[int, int, str]]) -> list[tuple[int, int, str]]:
    """spaCy 不接受重叠实体：按「更长者优先」保留。"""
    spans = sorted(spans, key=lambda s: (s[0] - s[1], s[0]))
    kept: list[tuple[int, int, str]] = []
    for s in spans:
        if all(s[1] <= k[0] or s[0] >= k[1] for k in kept):
            kept.append(s)
    return sorted(kept, key=lambda s: s[0])


# --------------------------------------------------------------------------
# 训练
# --------------------------------------------------------------------------


def train(data: list[tuple[str, dict]], out_dir: Path,
          base: str = "zh_core_web_sm", iterations: int = 30,
          dropout: float = 0.2, seed: int = 42) -> None:
    import spacy
    from spacy.training import Example
    from spacy.util import minibatch, compounding

    random.seed(seed)
    spacy.util.fix_random_seed(seed)

    # 起点模型：有通用中文模型就在其上微调，否则从空白中文管线开始
    try:
        nlp = spacy.load(base)
        print(f"[起点] 在 {base} 上微调")
    except Exception:
        nlp = spacy.blank("zh")
        print("[起点] 未找到通用中文模型，从空白 zh 管线训练")

    if "ner" not in nlp.pipe_names:
        ner = nlp.add_pipe("ner", last=True)
    else:
        ner = nlp.get_pipe("ner")

    for _, ann in data:
        for _s, _e, label in ann["entities"]:
            ner.add_label(label)

    # 只训练 NER，冻结其余组件
    other = [p for p in nlp.pipe_names if p != "ner"]
    with nlp.disable_pipes(*other):
        optimizer = (nlp.create_optimizer() if nlp.pipe_names == ["ner"]
                     else nlp.resume_training())
        for it in range(1, iterations + 1):
            random.shuffle(data)
            losses: dict = {}
            batches = minibatch(data, size=compounding(4.0, 32.0, 1.001))
            for batch in batches:
                examples = []
                for text, ann in batch:
                    try:
                        examples.append(
                            Example.from_dict(nlp.make_doc(text), ann)
                        )
                    except Exception:
                        continue
                if examples:
                    nlp.update(examples, drop=dropout, losses=losses,
                               sgd=optimizer)
            print(f"  iter {it:>3}/{iterations}  loss={losses.get('ner', 0):.2f}")

    out_dir.mkdir(parents=True, exist_ok=True)
    nlp.to_disk(out_dir)
    print(f"\n[完成] 模型已保存到 {out_dir}")
    print("      程序会自动发现该目录，把 NER 后端设为 auto / spacy 即可生效。")


# --------------------------------------------------------------------------


def main() -> int:
    ap = argparse.ArgumentParser(
        description="用弱监督方式训练通信工程领域 spaCy NER 模型",
    )
    ap.add_argument("--corpus", required=True, type=Path,
                    help="语料目录（.docx / .txt / .md）")
    ap.add_argument("--out", type=Path, default=_ROOT / "models" / "telecom_ner",
                    help="模型输出目录，默认 models/telecom_ner")
    ap.add_argument("--base", default="zh_core_web_sm",
                    help="微调起点模型，默认 zh_core_web_sm")
    ap.add_argument("--iters", type=int, default=30, help="训练轮数")
    ap.add_argument("--dump", type=Path, default=None,
                    help="可选：把自动标注结果导出为 JSONL 以便人工校对")
    args = ap.parse_args()

    if not args.corpus.is_dir():
        print(f"语料目录不存在：{args.corpus}", file=sys.stderr)
        return 2

    print(f"[1/3] 读取语料 {args.corpus} …")
    chunks = load_corpus(args.corpus)
    print(f"      得到 {len(chunks)} 个文本片段")
    if not chunks:
        print("语料为空，请放入 .docx / .txt 文件", file=sys.stderr)
        return 2

    print("[2/3] 用构词法规则自动标注 …")
    data = auto_annotate(chunks)
    total_ents = sum(len(a["entities"]) for _, a in data)
    print(f"      标注 {len(data)} 条样本 / {total_ents} 个实体")
    if total_ents < 50:
        print("      ⚠ 实体太少，训练效果会很差。建议扩充语料到数十份文档。")

    if args.dump:
        import json
        args.dump.parent.mkdir(parents=True, exist_ok=True)
        with open(args.dump, "w", encoding="utf-8") as f:
            for text, ann in data:
                f.write(json.dumps({"text": text, **ann},
                                   ensure_ascii=False) + "\n")
        print(f"      标注结果已导出：{args.dump}")

    if not data:
        return 2

    print(f"[3/3] 训练（{args.iters} 轮）…")
    train(data, args.out, base=args.base, iterations=args.iters)
    return 0


if __name__ == "__main__":
    sys.exit(main())
